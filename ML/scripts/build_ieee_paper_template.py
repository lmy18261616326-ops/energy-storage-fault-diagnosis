#!/usr/bin/env python3
"""Build the editable IEEE paper template from the verified project results."""

from __future__ import annotations

import re
import shutil
from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "ML" / "reports" / "ieee_paper_template_2026-08-04"
FIG_DIR = OUT / "figures"

TITLE = (
    "Physics-Guided and Observability-Aware Fault Diagnosis for a Bidirectional "
    "DC-DC Energy-Storage Converter: Simulink Modeling and Grouped Machine Learning"
)

AUTHORS = [
    "Mingyu Li",
    "Qiang Wang",
    "Xiaojing Xiong",
    "Bianjia Wang",
    "Zhongxin Cheng",
    "Ying Jia",
]

AFFILIATION = (
    "[VERIFY: School/Department], The Chinese University of Hong Kong, Shenzhen, "
    "Shenzhen 518172, China"
)

AUTHOR_NOTE = (
    "Qiang Wang, Xiaojing Xiong, and Bianjia Wang contributed equally and share "
    "second authorship. Zhongxin Cheng and Ying Jia are the corresponding authors. "
    "[VERIFY: insert the two corresponding-author e-mail addresses before submission.]"
)

DRAFT_NOTICE = (
    "EDITORIAL STATUS - COMPLETE SAMPLE MANUSCRIPT FOR AUTHOR REVISION. This paper is "
    "deliberately written as a simulation-only and computer-evidence study. It does not "
    "claim hardware, HIL, or laboratory validation. Every yellow [VERIFY: ...] item must "
    "be resolved by the authors before submission."
)

ABSTRACT = (
    "Reliable diagnosis of sensor and semiconductor-switch faults is necessary for "
    "bidirectional DC-DC converters used in battery energy-storage systems, yet model "
    "accuracy alone cannot compensate for faults that are weakly observable in the "
    "available signals. This paper presents a physics-guided and observability-aware "
    "diagnostic workflow built around a bidirectional buck-boost converter in MATLAB/"
    "Simulink. The main five-class study contains 416 independent simulation runs over "
    "28 operating points and uses 660 event-level features. Six-fold grouped out-of-fold "
    "validation keeps each operating point within a single fold. Logistic regression, "
    "random forest, and extremely randomized trees each attained a run-level macro-F1 "
    "and accuracy of 1.000; the exact 95% lower bound on accuracy was 0.991, while the "
    "healthy false-alarm rate was 0 with a 95% upper bound of 0.025. The feature table "
    "was strongly structured and low rank (effective rank 9.76), explaining why a "
    "regularized linear classifier matched tree ensembles and outperformed a 1D-CNN. "
    "For high-resistance switch degradation, generic controller measurements were "
    "insufficient. Adding synchronized device voltage and current enabled a physics "
    "threshold at 10.5 mOhm. An independent static set at an unseen 50 mOhm level was "
    "fully separated, and a 4320-run measurement-chain stress projection detected all "
    "12 mOhm cases under moderate and harsh error scenarios, with zero pre-fault false "
    "alarms; faults at or below 8 mOhm were not guaranteed. A supplemental 162-run "
    "synthetic transfer study exposed substantial domain shift and a possible seed/phase "
    "shortcut, thereby defining the claim boundary. The results indicate that converter "
    "observability, leakage-resistant grouped validation, and simple calibrated models "
    "are more valuable for this task than increasing network depth."
)

KEYWORDS = (
    "bidirectional DC-DC converter; battery energy storage; fault diagnosis; Simulink; "
    "observability; grouped cross-validation; high-resistance fault; machine learning"
)


def P(text: str) -> dict:
    return {"type": "p", "text": text}


def H(level: int, text: str) -> dict:
    return {"type": "h", "level": level, "text": text}


def B(items: list[str]) -> dict:
    return {"type": "bullets", "items": items}


def EQ(word: str, latex: str, label: str = "") -> dict:
    return {"type": "equation", "word": word, "latex": latex, "label": label}


def T(caption: str, columns: list[str], rows: list[list[str]], note: str = "") -> dict:
    return {"type": "table", "caption": caption, "columns": columns, "rows": rows, "note": note}


def F(filename: str, caption: str, source: str = "") -> dict:
    return {"type": "figure", "filename": filename, "caption": caption, "source": source}


def PLACEHOLDER(caption: str, instruction: str) -> dict:
    return {"type": "placeholder", "caption": caption, "instruction": instruction}


ELEMENTS = [
    H(1, "I. INTRODUCTION"),
    P(
        "Bidirectional DC-DC converters provide the controlled energy path between a "
        "battery and a DC bus. Their charge and discharge modes, large current stress, "
        "switching transitions, and tightly coupled control loops make them central to "
        "the dynamic performance of battery energy-storage systems. State-space "
        "averaging remains a standard basis for converter-oriented modeling {cite:middlebrook1976}, "
        "while recent bidirectional topologies seek lower ripple and higher efficiency "
        "{cite:lee2019}. Reviews and microgrid studies show that converter topology, "
        "control design, and operating mode must be treated together rather than as "
        "independent modules {cite:kondrath2017,sofla2011,xia2018}."
    ),
    P(
        "A converter fault can interrupt energy transfer or distort the measurements used "
        "by the controller. Existing work includes topology reconfiguration after "
        "open-circuit switch faults {cite:mahdavi2023}, observer-based open-switch diagnosis "
        "{cite:ding2024}, statistical residual testing for sensor faults {cite:alsheikh2015}, "
        "and sensor-fault-tolerant control based on unknown-input observers "
        "{cite:ouahabi2025}. These methods establish that physical residuals and model "
        "structure are valuable. However, a data-driven classifier can still report "
        "optimistic scores if windows from the same operating condition are split between "
        "training and testing, if injected labels are embedded in the feature table, or if "
        "the selected sensors do not physically expose the target degradation."
    ),
    P(
        "Battery diagnosis research increasingly combines model-based and data-driven "
        "evidence {cite:jin2024}, and recent reviews identify sensor fusion, artificial "
        "intelligence, uncertainty, and transferable validation as key topics "
        "{cite:liu2025,xu2025}. In a converter study, these ideas require an additional "
        "constraint: the diagnostic channel must be observable at the electrical location "
        "where the fault changes the circuit. The present work therefore treats model "
        "selection as the last stage of a chain that begins with the Simulink model, fault "
        "injection, and measurement design."
    ),
    P("The main contributions are summarized as follows."),
    B([
        "An end-to-end Simulink-to-machine-learning workflow is formulated for five run-level classes: healthy operation, DC-bus voltage-sensor bias, inductor-current-sensor bias, and two switch-side fault families.",
        "A leakage-resistant six-fold grouped out-of-fold protocol is used across 28 operating points. Exact binomial confidence bounds and healthy false-alarm bounds are reported together with point estimates.",
        "The geometry of the 660-feature event table is analyzed. Its low effective rank explains why regularized logistic regression and tree ensembles are suitable, whereas a 1D-CNN is mismatched to the available sample size and representation.",
        "A negative high-resistance result is used to redesign observability. Synchronized device voltage and current support a physics threshold, followed by measurement-error and timing-jitter stress projections.",
        "A separate synthetic waveform domain is used as a transfer and shortcut audit. Its failure modes are retained in the paper rather than hidden, defining what the current evidence can and cannot support.",
    ]),
    P(
        "The manuscript is organized as follows. Section II presents the converter and "
        "control model. Section III defines faults and observability. Sections IV and V "
        "describe data construction, features, candidate models, and validation. Section "
        "VI reports the comparative results. Sections VII and VIII discuss implications, "
        "limitations, and the simulation-only claim boundary."
    ),

    H(1, "II. BIDIRECTIONAL CONVERTER AND SIMULINK MODEL"),
    H(2, "A. System Architecture"),
    P(
        "The study considers a non-isolated synchronous bidirectional buck-boost stage "
        "connecting a low-voltage battery port to a regulated DC bus. In discharge, the "
        "converter operates in boost mode and transfers battery power to the bus; in "
        "charge, it operates in buck mode and absorbs bus power into the battery. This "
        "choice is representative of storage interfaces described in {cite:kondrath2017,sofla2011,xia2018}; "
        "multi-port storage converters provide a related path toward integrated battery "
        "management and module bypass {cite:alsonisi2024}. "
        "The final submitted paper should include one schematic exported directly from the "
        "validated Simulink model so that switch naming, current polarity, and sensing "
        "locations exactly match the equations and fault labels."
    ),
    PLACEHOLDER(
        "Fig. 1. Bidirectional buck-boost power stage, control loops, and sensing locations used in the final Simulink model.",
        "REPLACE THIS BOX with a 600-dpi vector or PNG export from main_model_fd_v06_switchobservability.slx. Label S1, S2, L, battery voltage/current, DC-bus voltage, device voltage, device current, duty command, and charge/discharge current polarity."
    ),
    P(
        "For a compact description that is valid across the two modes, define the state "
        "vector x = [iL, vdc]^T, input vector u = [vbat, iload]^T, duty ratio d, operating "
        "mode m, parameter vector theta, physical fault fp, sensor fault fs, and measurement "
        "noise n. The averaged switched model is written as"
    ),
    EQ(
        "x_dot = A_m(d, theta) x + B_m(d, theta) u + E_m f_p,    y = C_m x + D_m u + f_s + n.",
        r"\dot{\mathbf{x}}=\mathbf{A}_{m}(d,\boldsymbol{\theta})\mathbf{x}+\mathbf{B}_{m}(d,\boldsymbol{\theta})\mathbf{u}+\mathbf{E}_{m}\mathbf{f}_{p},\quad \mathbf{y}=\mathbf{C}_{m}\mathbf{x}+\mathbf{D}_{m}\mathbf{u}+\mathbf{f}_{s}+\mathbf{n}.",
        "(1)",
    ),
    P(
        "Equation (1) separates changes in electrical dynamics from additive or gain-type "
        "sensor faults. It also permits the exact switching implementation to remain in "
        "Simulink while the paper uses a mode-dependent averaged representation for "
        "analysis. [VERIFY: insert the final A_m and B_m matrices, or an appendix derivation, "
        "after confirming the duty-ratio convention of the submitted model.]"
    ),
    H(2, "B. Cascaded Control"),
    P(
        "The control architecture contains an outer voltage or power loop and an inner "
        "inductor-current loop. With bus-voltage error ev = vdc* - vdc and current error "
        "ei = iL* - iL, a generic implementation is"
    ),
    EQ(
        "iL* = Kpv ev + Kiv integral(ev dt),    d = sat[dff + Kpi ei + Kii integral(ei dt)].",
        r"i_L^{\star}=K_{pv}e_v+K_{iv}\int e_v\,dt,\quad d=\operatorname{sat}\!\left[d_{ff}+K_{pi}e_i+K_{ii}\int e_i\,dt\right].",
        "(2)",
    ),
    P(
        "Mode supervision selects the sign of the current reference and the active PWM "
        "logic. Saturation, anti-windup, current limiting, and protection states must be "
        "represented because their activation changes the waveform distribution and may "
        "be mistaken for a fault by a classifier. In the present pipeline, controller "
        "commands, measured electrical variables, and protection flags are logged with a "
        "run identifier and an operating-point identifier."
    ),
    T(
        "TABLE I. PRINCIPAL SIMULATION AND DIAGNOSTIC SETTINGS",
        ["Item", "Value used in current evidence", "Submission action"],
        [
            ["Converter", "Non-isolated synchronous bidirectional buck-boost", "Confirm final schematic and polarity"],
            ["Nominal DC bus", "400 V nominal", "[VERIFY: final mask value and tolerance]"],
            ["Switching frequency", "20 kHz (50 us PWM period)", "Confirmed by cycle aggregation setting"],
            ["Device-channel sample interval", "1 us in high-resistance studies", "State anti-aliasing and bandwidth assumptions"],
            ["Main operating points", "28", "List voltage, current, load, and mode grid in appendix"],
            ["Main independent runs", "416", "Retain immutable run IDs in released data"],
            ["High-resistance threshold", "10.5 mOhm equivalent resistance", "Do not claim guaranteed detection below 12 mOhm"],
            ["Hardware/HIL", "Not performed", "Keep limitation explicit in abstract and conclusion"],
        ],
        "Values marked [VERIFY] are intentionally not inferred from filenames or plotting conventions.",
    ),
    H(2, "C. Reproducibility Boundary"),
    P(
        "The converter model, injection logic, logging configuration, feature code, and "
        "fold assignment must be versioned as a single experiment. A model result is "
        "reproducible only when the Simulink revision and dataset manifest are frozen. The "
        "current evidence uses recorded file hashes, modification times, analysis "
        "parameters, and random seed 20260804 for the measurement-chain stress projection."
    ),

    H(1, "III. FAULT MODELING AND OBSERVABILITY-AWARE DIAGNOSIS"),
    H(2, "A. Fault Taxonomy"),
    P(
        "The main classifier uses five run-level labels. Healthy runs include nominal and "
        "targeted healthy variations. Two sensor classes represent bias at the DC-bus "
        "voltage and inductor-current channels. The remaining labels identify whether the "
        "switching abnormality is associated with the S1 or S2 side; partial-open, full-open, "
        "and intermittent injections are treated as members of the corresponding side "
        "family for the principal comparison. This family-level formulation avoids "
        "claiming subtype resolution when the number of independent runs is limited."
    ),
    T(
        "TABLE II. FAULT CLASSES, INJECTION FORM, AND PRIMARY OBSERVABILITY",
        ["Class", "Run support", "Injection concept", "Primary evidence"],
        [
            ["Healthy", "144", "Nominal and targeted healthy variation", "Stable control and physical residuals"],
            ["DC-bus voltage-sensor bias", "80", "Additive or gain bias in vdc measurement", "Voltage reference error and cross-channel consistency"],
            ["Inductor-current-sensor bias", "80", "Additive or gain bias in iL measurement", "Current/power residuals and controller response"],
            ["S1-side switch-fault family", "56", "Partial, full-open, or intermittent switching abnormality", "Mode-aware waveform/event features"],
            ["S2-side switch-fault family", "56", "Partial, full-open, or intermittent switching abnormality", "Mode-aware waveform/event features"],
            ["High-resistance degradation", "Separate study", "Increased conducting-path resistance", "Synchronized device voltage/current ratio"],
        ],
        "The five-class total is 416 independent runs. High resistance is evaluated separately because it requires a different sensor set and decision rule.",
    ),
    H(2, "B. Sensor Faults and Switch Faults"),
    P(
        "A sensor channel can be represented as ym = (1 + g)y + b + n, where g is gain "
        "error and b is offset. The resulting fault signature depends on loop location: a "
        "biased feedback variable changes both the recorded signal and the duty command. "
        "Therefore, features include cross-channel and command-response consistency rather "
        "than the suspect sensor alone. Open and intermittent switch faults alter the "
        "relationship between gate command, inductor current, bus voltage, and delivered "
        "power. Observer and residual methods in {cite:ding2024,alsheikh2015,ouahabi2025} "
        "motivate this use of redundant physical relations."
    ),
    H(2, "C. High-Resistance Observability"),
    P(
        "The first high-resistance attempt used generic controller-level measurements. On "
        "16 independent runs it produced fault recall of 0.625 and healthy false-alarm "
        "rate of 0.25; none of 28 candidate signal formulations met the acceptance rule. "
        "This negative result showed that model complexity was not the main bottleneck. "
        "The revised v06 Simulink model therefore logs synchronized voltage across and "
        "current through the relevant conducting device."
    ),
    P("A physically interpretable equivalent-resistance residual is then defined as"),
    EQ(
        "rR(t) = |vdevice(t)| / (|idevice(t)| + epsilon),",
        r"r_R(t)=\frac{|v_{\mathrm{device}}(t)|}{|i_{\mathrm{device}}(t)|+\varepsilon},",
        "(3)",
    ),
    P(
        "with conduction gating and a small epsilon to avoid division at near-zero current. "
        "Because the 1 us device samples and 50 us PWM period share a fixed timing relation, "
        "the residual is aggregated within complete PWM cycles before a run decision. This "
        "prevents a fixed switching phase from dominating the estimate. The frozen threshold "
        "is 0.0105 ohm, derived from the separation between the development healthy maximum "
        "(0.001 ohm) and development fault minimum (0.020 ohm)."
    ),
    F(
        "diagnostic_architecture.png",
        "Fig. 2. Physics-guided and observability-aware diagnostic workflow used in this study.",
        "Generated from the implemented analysis stages; the final paper should align signal names with Fig. 1.",
    ),

    H(1, "IV. DATASET CONSTRUCTION AND FEATURE ENGINEERING"),
    H(2, "A. Main Five-Class Simulation Dataset"),
    P(
        "The principal dataset contains 416 independent runs across 28 operating points. "
        "Each run has a unique identifier, one operating-point group, and one run-level "
        "label. Operating points vary the commanded mode and electrical loading conditions "
        "[VERIFY: provide the complete voltage-current-load grid in a supplementary table]. "
        "The grouping unit is the operating point rather than an arbitrary time window."
    ),
    P(
        "Event extraction aligns pre-fault and post-fault behavior without using the true "
        "fault value, injection code, true device state, or any label-derived field as a "
        "model input. Candidate feature families include robust location and dispersion, "
        "slopes, differences between pre-event and post-event segments, duty/current "
        "coupling, voltage-current-power consistency, and mode-aware protection responses. "
        "The final table has 660 numeric features."
    ),
    H(2, "B. Feature Geometry and Leakage Controls"),
    P(
        "The event table is high dimensional by column count but not by intrinsic geometry. "
        "Twenty-six features are constant, 5803 feature pairs have absolute correlation at "
        "least 0.95, and only 10, 21, and 53 principal components are required for 90%, 95%, "
        "and 99% explained variance, respectively. The entropy-based effective rank is 9.758. "
        "These statistics imply a small number of strong physical directions repeated across "
        "many related summaries."
    ),
    F(
        "pca_geometry.png",
        "Fig. 3. Principal-component geometry of the main 660-feature event table.",
        "Copied from the verified model-selection analysis artifact.",
    ),
    P("To limit information leakage, the following exclusions are enforced."),
    B([
        "All windows and events from one operating point remain in one validation fold.",
        "Scaling, imputation, and any feature selection are fitted on the training portion of each fold only.",
        "Ground-truth resistance, actual switch state, injection time, activity/effectiveness flags, fault code, and labels are excluded from training inputs.",
        "Metrics are aggregated at the run level; correlated windows are not counted as independent confidence-interval samples.",
        "Duplicate prediction keys and required-field omissions are checked before metric computation.",
    ]),
    H(2, "C. Supplemental Synthetic Waveform Domain"),
    P(
        "A separately generated HDF5 dataset is used only as a supplemental transfer audit. "
        "It contains 162 runs, 12 operating points, 34 channels, and 20,001 samples per run "
        "at 1 us spacing. Its metadata explicitly states that it is synthetic engineering "
        "reference data and not measured data. Consequently, it cannot replace dynamic "
        "Simulink power-stage simulation or experimental evidence."
    ),
    P(
        "Two representations are evaluated: controller-observable channels and a "
        "device-augmented set that also includes S1/S2 device voltage and current. A pre-fault "
        "0-6 ms negative control is evaluated to detect label information encoded in random "
        "seed or waveform phase. The frozen main-domain model is also applied without "
        "retraining to quantify transfer failure."
    ),

    H(1, "V. CANDIDATE MODELS AND VALIDATION PROTOCOL"),
    H(2, "A. Model Families"),
    P(
        "Seven model families are compared under the same run and group assignments. "
        "Regularized multinomial logistic regression provides a low-variance linear "
        "baseline. Random forest and extremely randomized trees represent bagged nonlinear "
        "ensembles, while gradient-boosted trees model residual nonlinear structure. K-nearest "
        "neighbors tests local geometry after scaling. A multilayer perceptron tests a dense "
        "nonlinear mapping, and a compact 1D-CNN operates on down-sampled raw multichannel "
        "sequences."
    ),
    P(
        "The models are not ranked by architecture novelty. They are ranked by grouped "
        "out-of-fold macro-F1, accuracy, healthy false-alarm rate, uncertainty, training "
        "cost, and interpretability. A physics threshold is included for the high-resistance "
        "subproblem because the residual in (3) has a known electrical meaning."
    ),
    H(2, "B. Grouped Out-of-Fold Evaluation"),
    P(
        "The 28 operating points are partitioned into six folds. For each fold, all runs from "
        "the held-out operating points are predicted by a model trained on the remaining "
        "groups. Concatenating the six held-out predictions gives one out-of-fold prediction "
        "per independent run. No score reported for the main table is an in-sample score."
    ),
    P("For C classes and run-level confusion counts, macro-F1 is"),
    EQ(
        "Macro-F1 = (1/C) sum_c [2 Precision_c Recall_c / (Precision_c + Recall_c)].",
        r"\operatorname{MacroF1}=\frac{1}{C}\sum_{c=1}^{C}\frac{2P_cR_c}{P_c+R_c}.",
        "(4)",
    ),
    P(
        "Healthy false-alarm rate is the fraction of healthy runs assigned to any fault "
        "class. Exact two-sided 95% binomial bounds are used for accuracy and class recall; "
        "the one-sided upper bound is emphasized when zero healthy false alarms are observed. "
        "Confidence statements use independent runs, not the number of windows."
    ),
    H(2, "C. Model Selection and Fusion Rule"),
    P(
        "A more complex model is retained only when it improves an identified weakness on "
        "grouped held-out runs. On the main task, logistic regression, random forest, and "
        "extremely randomized trees make the same correct decisions. Their probability "
        "fusion therefore adds no observed accuracy benefit and would complicate calibration. "
        "The preferred deployable classifier is logistic regression, with a tree ensemble "
        "as a sensitivity model and the physics threshold as a separate high-resistance "
        "safety channel."
    ),

    H(1, "VI. RESULTS"),
    H(2, "A. Main Five-Class Comparison"),
    P(
        "Table III reports the 416-run grouped out-of-fold comparison. Logistic regression, "
        "random forest, and extremely randomized trees each classify every run correctly. "
        "The point estimate of 1.000 must not be interpreted as guaranteed field performance: "
        "the exact 95% lower bound on accuracy is 0.991. Among 144 healthy runs, no false "
        "alarm is observed, but the exact 95% upper bound remains 0.0253. For each 56-run "
        "switch-side family, recall is 1.000 with a 95% lower bound of 0.936."
    ),
    T(
        "TABLE III. SIX-FOLD GROUPED OOF RESULTS ON THE MAIN FIVE-CLASS DATASET",
        ["Model", "Macro-F1", "Accuracy", "95% accuracy interval", "Healthy FAR", "95% FAR upper"],
        [
            ["Logistic regression", "1.0000", "1.0000", "[0.9912, 1.0000]", "0.0000", "0.0253"],
            ["Random forest", "1.0000", "1.0000", "[0.9912, 1.0000]", "0.0000", "0.0253"],
            ["Extra Trees", "1.0000", "1.0000", "[0.9912, 1.0000]", "0.0000", "0.0253"],
            ["K-nearest neighbors", "0.9901", "0.9880", "[0.9722, 0.9961]", "0.0000", "0.0253"],
            ["Gradient-boosted trees", "0.9887", "0.9904", "[0.9756, 0.9974]", "0.0000", "0.0253"],
            ["Multilayer perceptron", "0.9763", "0.9712", "[0.9502, 0.9850]", "0.0000", "0.0253"],
            ["1D-CNN", "0.8466", "0.8125", "[0.7716, 0.8489]", "0.3194", "0.4022"],
        ],
        "OOF = out-of-fold; FAR = false-alarm rate. Confidence intervals use independent runs.",
    ),
    F(
        "main_model_run_metrics.png",
        "Fig. 4. Run-level grouped out-of-fold performance of the seven candidate models.",
        "Copied from the verified IEEE computer-evidence artifact.",
    ),
    P(
        "The 1D-CNN performs substantially worse, including a healthy false-alarm rate of "
        "0.319. This is consistent with the task geometry: only 416 independent runs are "
        "available, the engineered representation is already aligned to event physics, and "
        "the dominant class directions are low dimensional. Convolution over the raw "
        "sequence adds parameters and phase sensitivity without a demonstrated generalization "
        "benefit."
    ),
    H(2, "B. High-Resistance Static Separation"),
    P(
        "With direct device observability, the development set contains 32 runs and 3808 "
        "cycle-level windows. An independent set contains 16 runs and 1904 windows at an "
        "unseen 0.05 ohm fault level. The frozen 0.0105 ohm threshold and all tested "
        "comparators achieve run-level macro-F1 1.000, fault recall 1.000, and healthy FAR "
        "0. The physics threshold is preferred because it is interpretable and does not "
        "require a high-capacity classifier. This result demonstrates separability in the "
        "idealized v06 electrical model; it does not establish temperature-robust hardware "
        "sensitivity."
    ),
    H(2, "C. Measurement-Chain Stress Projection"),
    P(
        "The robustness projection reuses 24 independent healthy Simulink current trajectories. "
        "Nine resistance severities, four measurement scenarios, and five random repetitions "
        "produce 4320 projected runs. Voltage and current noise, gain error, offset, quantization, "
        "healthy-resistance drift, and 1-5 us channel misalignment are varied. This is a "
        "Monte Carlo projection of the sensing and decision chain, not a new dynamic power-stage "
        "simulation and not a hardware experiment."
    ),
    T(
        "TABLE IV. HIGH-RESISTANCE DETECTION BOUNDARY IN THE STRESS PROJECTION",
        ["Commanded resistance", "Moderate scenario", "Harsh scenario", "Interpretation"],
        [
            ["<= 8 mOhm", "0 detections", "0 detections", "Not guaranteed; early degradation region"],
            ["10 mOhm", "0/120", "8/120", "Near threshold and sensitive to error realization"],
            ["12 mOhm", "120/120", "120/120", "Current evidence-supported lower guarantee boundary"],
            [">= 15 mOhm", "All detected", "All detected", "Detected across evaluated scenarios"],
        ],
        "At 12 mOhm the exact 95% lower detection bound is 0.9697. Pre-fault FAR is 0 in the evaluated projection.",
    ),
    F(
        "high_r_detection_curve.png",
        "Fig. 5. Detection probability versus high-resistance severity under measurement-chain stress.",
        "The plot is a sensing-chain projection driven by real Simulink healthy trajectories, not hardware evidence.",
    ),
    F(
        "high_r_threshold_tradeoff.png",
        "Fig. 6. Frozen-threshold tradeoff for the equivalent-resistance residual.",
        "The selected threshold is 10.5 mOhm; the defensible guarantee region begins at 12 mOhm.",
    ),
    P(
        "At 12 mOhm, 120 of 120 moderate and 120 of 120 harsh cases are detected. The "
        "moderate-scenario median latency is 12.91 ms and the 95th percentile is 14.78 ms. "
        "At 10 mOhm, the moderate scenario yields no detections and the harsh scenario yields "
        "8 of 120, while severities at or below 8 mOhm are not detected. The correct claim is "
        "therefore a stress-tested 12 mOhm boundary, not universal early-degradation detection."
    ),
    H(2, "D. Supplemental Synthetic Transfer Audit"),
    P(
        "The frozen main-domain model degrades sharply on the independently generated "
        "synthetic HDF5 domain: macro-F1 is 0.4748, accuracy is 0.5417, and healthy FAR is "
        "0.3333 within the applicable active scope. This is evidence of domain shift rather "
        "than a reason to retune the reported main result after seeing the target labels."
    ),
    T(
        "TABLE V. RETRAINED SIX-CLASS GROUPED OOF MACRO-F1 ON THE SYNTHETIC TRANSFER DOMAIN",
        ["Model", "Controller-observable", "Device-augmented"],
        [
            ["Logistic regression", "0.8023", "0.9230"],
            ["Random forest", "0.9149", "0.9930"],
            ["Extra Trees", "0.9493", "1.0000"],
            ["Gradient-boosted trees", "0.9376", "1.0000"],
            ["K-nearest neighbors", "0.7318", "0.9252"],
            ["Multilayer perceptron", "0.7465", "0.9074"],
            ["1D-CNN", "0.3822", "0.6201"],
        ],
        "These data are explicitly synthetic, not measured. Values are supportive sensitivity results and are not the main validation claim.",
    ),
    F(
        "synthetic_model_comparison.png",
        "Fig. 7. Model comparison in the supplemental synthetic waveform domain.",
        "The device-augmented result is potentially affected by synthetic generation shortcuts and must not be presented as independent experimental validation.",
    ),
    P(
        "Retraining shows that direct device channels substantially improve most tabular "
        "models. However, the best pre-fault negative-control macro-F1 is 0.2319, above the "
        "six-class chance level of 0.1667. Inspection of the generator shows that nonoverlapping "
        "random-seed ranges and an 850 Hz phase term can carry label-related information. "
        "Therefore, the apparently perfect device-augmented Extra Trees result is retained "
        "only as a sensitivity demonstration. A future blind dataset must use paired seeds "
        "shared across healthy and fault labels or make all nuisance seeds label independent."
    ),

    H(1, "VII. DISCUSSION"),
    H(2, "A. Why Simple Models Are Most Suitable"),
    P(
        "The main event table is a small-sample, high-column-count, low-effective-rank "
        "distribution. Many columns are alternate summaries of the same control and energy "
        "relations, so class separation occurs along a small number of physical directions. "
        "A regularized linear decision boundary is therefore plausible, while tree ensembles "
        "provide a check for nonlinear thresholds and interactions. The identical top-level "
        "run decisions of logistic regression, random forest, and Extra Trees indicate that "
        "additional model capacity is unnecessary on the current scope."
    ),
    P(
        "The preferred main classifier is logistic regression because it matches the best "
        "grouped OOF accuracy, is inexpensive to train, is straightforward to calibrate, and "
        "supports coefficient-level interpretation. Extra Trees should be retained as a "
        "secondary sensitivity model. The high-resistance branch should remain a separate "
        "physics residual because its input requirements and detection boundary differ from "
        "the five-class event classifier."
    ),
    H(2, "B. Why Model Fusion Does Not Improve the Main Task"),
    P(
        "Fusion is useful when constituent models make complementary errors. Here, the three "
        "best tabular models make no run-level errors on the evaluated main data, and their "
        "error sets are therefore not complementary. Probability averaging can change "
        "confidence without adding evidence for improved decisions. A practical architecture "
        "is instead modular: a calibrated logistic classifier for common active faults, a "
        "tree model for audit, and a device-level resistance residual for gradual conduction "
        "degradation."
    ),
    H(2, "C. Observability Is More Important Than Classifier Capacity"),
    P(
        "The high-resistance study provides the clearest causal lesson. With generic control "
        "signals, 28 attempted formulations did not satisfy recall and false-alarm targets. "
        "After adding the voltage and current that define conducting-path resistance, a simple "
        "threshold separated the idealized model. The improvement came from measuring the "
        "fault mechanism, not from changing the learning algorithm. This supports an "
        "observability-first workflow for converter fault diagnosis."
    ),
    H(2, "D. Relation to Existing Work"),
    P(
        "Observer-based methods provide physical residuals and can offer analytical guarantees "
        "under stated modeling assumptions {cite:ding2024,alsheikh2015,ouahabi2025}. The "
        "present approach is complementary: Simulink defines the physical system and fault "
        "injection, event features summarize redundant measurements, and grouped learning "
        "captures mode-dependent decision surfaces. The separate high-resistance threshold "
        "preserves a physical interpretation where the relevant voltage-current pair is "
        "available. Hybrid model/data logic is consistent with broader battery diagnosis "
        "developments {cite:jin2024,liu2025,xu2025}."
    ),
    H(2, "E. Practical Interpretation for an IEEE Simulation Paper"),
    P(
        "The current contribution is strongest as a methodological simulation paper whose "
        "novelty is the integration of fault observability, leakage-resistant validation, "
        "model-selection evidence, and explicit claim boundaries. The main result is not that "
        "a complex classifier achieves a headline score. It is that the diagnostic design "
        "explains when simple models are sufficient and when sensing must be changed. This "
        "position is appropriate for a general IEEE conference or a simulation-oriented "
        "journal submission, provided that the final manuscript includes the complete "
        "Simulink topology, parameters, injection schedule, and reproducible dataset manifest."
    ),

    H(1, "VIII. LIMITATIONS AND THREATS TO VALIDITY"),
    P(
        "First, no hardware, controller-HIL, or power-HIL experiment is included. The reported "
        "metrics establish behavior within the modeled and projected domains only. Component "
        "parasitics, sensor common-mode behavior, switching-node bandwidth, electromagnetic "
        "interference, timing implementation, and protection interactions may reduce field "
        "performance."
    ),
    P(
        "Second, the high-resistance Simulink model uses an idealized conducting resistance. "
        "It does not yet represent temperature-dependent semiconductor voltage drop, bond-wire "
        "or solder degradation, junction-temperature feedback, or aging-dependent packaging. "
        "The 4320-run study projects measurement-chain variation onto healthy dynamic current "
        "trajectories; it is not a dynamic high-resistance power-stage simulation."
    ),
    P(
        "Third, confidence intervals are limited by the number of independent runs. The 1904 "
        "high-resistance windows in the independent static set are correlated within 16 runs "
        "and cannot be treated as 1904 independent samples. The lower recall bounds for the "
        "56-run switch families and the upper FAR bound for 144 healthy runs should remain "
        "visible."
    ),
    P(
        "Fourth, the synthetic HDF5 domain is not measured data and contains a plausible "
        "seed/phase shortcut. It is useful for detecting domain shift and feature dependence, "
        "but not for a blind generalization claim. Finally, the main five-class label merges "
        "several switch-fault subtypes by side; subtype-level performance requires more "
        "independent runs and should be reported only after a separate grouped study."
    ),
    P(
        "These limitations are not post hoc qualifications. They define the intended use of "
        "the current paper: a reproducible simulation and model-selection study, with hardware "
        "validation left as future work rather than implied by synthetic evidence."
    ),

    H(1, "IX. CONCLUSION"),
    P(
        "This paper developed a physics-guided and observability-aware diagnostic workflow for "
        "a bidirectional DC-DC battery energy-storage converter modeled in Simulink. On 416 "
        "independent runs across 28 operating points, six-fold grouped out-of-fold validation "
        "showed that logistic regression, random forest, and Extra Trees achieved perfect "
        "point estimates for the five-class active-observable scope, with an exact 95% accuracy "
        "lower bound of 0.991. Data geometry showed that the 660-feature table had effective "
        "rank 9.76, supporting the choice of simple regularized and tree models over a 1D-CNN."
    ),
    P(
        "For gradual high-resistance degradation, controller-level signals were inadequate. "
        "Direct synchronized device voltage and current enabled a 10.5 mOhm physics threshold. "
        "The stress projection supported complete detection at 12 mOhm under the evaluated "
        "moderate and harsh measurement conditions, while faults at or below 8 mOhm remained "
        "outside the guaranteed region. A synthetic transfer audit further demonstrated that "
        "strong within-domain performance does not imply cross-domain validity."
    ),
    P(
        "The central conclusion is that observability and validation design dominate classifier "
        "complexity for the studied converter. Future work should replace the ideal resistance "
        "with temperature-dependent semiconductor and interconnect degradation, generate a "
        "label-independent blind synthetic set, and, when resources permit, evaluate the frozen "
        "pipeline on HIL or hardware. The present claims remain explicitly simulation-only."
    ),

    H(1, "DATA AND CODE AVAILABILITY"),
    P(
        "The manuscript was prepared from versioned Simulink models, run manifests, feature "
        "tables, grouped out-of-fold predictions, confidence-interval summaries, and plotting "
        "scripts stored in the project workspace. [VERIFY: insert the public repository or "
        "institutional archive DOI before submission.] The supplemental HDF5 data must be "
        "distributed with its original statement: synthetic engineering reference data; NOT "
        "measured data."
    ),

    H(1, "AUTHOR CONTRIBUTIONS"),
    P(
        "[VERIFY WITH ALL AUTHORS BEFORE SUBMISSION.] Mingyu Li: conceptualization, methodology, "
        "software, validation, formal analysis, visualization, and writing - original draft. "
        "Qiang Wang, Xiaojing Xiong, and Bianjia Wang: investigation, data curation, validation, "
        "and writing - review and editing. Zhongxin Cheng and Ying Jia: supervision, project "
        "administration, funding acquisition, and writing - review and editing."
    ),

    H(1, "FUNDING AND ACKNOWLEDGMENT"),
    P(
        "Funding statement template - verify before submission: This work was supported by the "
        "University Development Fund of The Chinese University of Hong Kong, Shenzhen under "
        "Project No. [VERIFY: INSERT THE OFFICIAL PROJECT/GRANT NUMBER]. The authors acknowledge "
        "the computational resources provided by The Chinese University of Hong Kong, Shenzhen. "
        "[VERIFY: retain only statements supported by the actual award and institutional wording.]"
    ),

    H(1, "CONFLICT OF INTEREST"),
    P(
        "The authors declare no conflict of interest. [VERIFY: confirm this statement with all "
        "authors and disclose any relevant financial or personal relationships.]"
    ),
]


REFERENCES = [
    (
        "middlebrook1976",
        "R. D. Middlebrook and S. Cuk, \"A general unified approach to modelling switching-converter power stages,\" in Proc. IEEE Power Electronics Specialists Conf. (PESC), 1976, pp. 18-34, doi: 10.1109/PESC.1976.7072895.",
    ),
    (
        "lee2019",
        "H.-S. Lee and J.-J. Yun, \"High-efficiency bidirectional buck-boost converter for photovoltaic and energy storage systems in a smart grid,\" IEEE Trans. Power Electron., vol. 34, no. 5, pp. 4316-4328, May 2019, doi: 10.1109/TPEL.2018.2860059.",
    ),
    (
        "kondrath2017",
        "N. Kondrath, \"Bidirectional DC-DC converter topologies and control strategies for interfacing energy storage systems in microgrids: An overview,\" in Proc. 5th IEEE Int. Conf. Smart Energy Grid Engineering (SEGE), 2017, pp. 341-345.",
    ),
    (
        "sofla2011",
        "M. A. Sofla and L. Wang, \"Control of DC-DC bidirectional converters for interfacing batteries in microgrids,\" in Proc. IEEE/PES Power Systems Conf. Expo. (PSCE), 2011, pp. 1-6, doi: 10.1109/PSCE.2011.5772602.",
    ),
    (
        "xia2018",
        "K. Xia, X. Hong, Y. Yuan, and A. Chen, \"Bi-directional DC-DC converter for battery energy storage,\" J. System Simulation, vol. 30, no. 8, pp. 3219-3228, Aug. 2018, doi: 10.16182/j.issn1004731x.joss.201808049, in Chinese.",
    ),
    (
        "mahdavi2023",
        "M. S. Mahdavi, M. S. Karimzadeh, T. Rahimi, and G. B. Gharehpetian, \"A fault-tolerant bidirectional converter for battery energy storage systems in DC microgrids,\" Electronics, vol. 12, no. 3, Art. no. 679, 2023, doi: 10.3390/electronics12030679.",
    ),
    (
        "ding2024",
        "S. Ding, D. Tang, J. Hang, J. Zhao, and S. Gui, \"Robust open-switch fault diagnosis of bidirectional DC/DC converters based on extended Kalman filter with multiple corrections,\" IEEE Trans. Circuits Syst. I, Reg. Papers, vol. 71, no. 9, pp. 4363-4374, Sep. 2024, doi: 10.1109/TCSI.2024.3421655.",
    ),
    (
        "alsheikh2015",
        "H. Al-Sheikh, G. Hoblos, N. Moubayed, and N. Karami, \"A sensor fault diagnosis scheme for a DC/DC converter used in hybrid electric vehicles,\" IFAC-PapersOnLine, vol. 48, no. 21, pp. 713-719, 2015, doi: 10.1016/j.ifacol.2015.09.611.",
    ),
    (
        "ouahabi2025",
        "M. S. Ouahabi et al., \"Real-time sensor fault tolerant control of DC-DC converters in DC microgrids using a switching unknown input observer,\" IEEE Access, vol. 13, pp. 95838-95850, 2025, doi: 10.1109/ACCESS.2025.3571650.",
    ),
    (
        "jin2024",
        "H. Jin, Z. Gao, Z. Zuo, Z. Zhang, Y. Wang, and A. Zhang, \"A combined model-based and data-driven fault diagnosis scheme for lithium-ion batteries,\" IEEE Trans. Ind. Electron., vol. 71, no. 6, pp. 6274-6284, Jun. 2024, doi: 10.1109/TIE.2023.3299029.",
    ),
    (
        "liu2025",
        "K. Liu, S. Zhao, Y. Wang, K. Li, J. Wang, Y. Sun, Q. Wu, and Q. Peng, \"Advanced fault diagnosis in batteries: Insights into fault mechanisms, sensor fusion, and artificial intelligence,\" Adv. Appl. Energy, vol. 20, Art. no. 100247, 2025, doi: 10.1016/j.adapen.2025.100247.",
    ),
    (
        "xu2025",
        "Y. Xu, X. Ge, R. Guo, and W. Shen, \"Recent advances in model-based fault diagnosis for lithium-ion batteries: A comprehensive review,\" Renew. Sustain. Energy Rev., vol. 207, Art. no. 114922, 2025, doi: 10.1016/j.rser.2024.114922.",
    ),
    (
        "alsonisi2024",
        "M. Alsonisi, S. Ethni, M. Elgendy, M. Ahmeid, and B. Yildirim, \"DC-DC bidirectional converter for battery energy storage system with integrated battery management,\" in Proc. IEEE Int. Conf. Expo. Electric Power Engineering (EPEi), 2024, pp. 89-93, doi: 10.1109/EPEi63510.2024.10758159.",
    ),
]

REF_NUM = {key: idx + 1 for idx, (key, _) in enumerate(REFERENCES)}


def citation_word(text: str) -> str:
    def repl(match: re.Match[str]) -> str:
        nums = [str(REF_NUM[key.strip()]) for key in match.group(1).split(",")]
        return "[" + "], [".join(nums) + "]"

    return re.sub(r"\{cite:([^}]+)\}", repl, text)


def latex_escape_plain(text: str) -> str:
    cite_tokens: list[str] = []

    def save_cite(match: re.Match[str]) -> str:
        cite_tokens.append(match.group(1))
        return f"CITETOKEN{len(cite_tokens)-1}END"

    text = re.sub(r"\{cite:([^}]+)\}", save_cite, text)
    text = text.replace("µ", "MICROTOKEN").replace("Ω", "OHMTOKEN")
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    text = "".join(replacements.get(ch, ch) for ch in text)
    text = text.replace("MICROTOKEN", r"$\mu$").replace("OHMTOKEN", r"$\Omega$")
    for idx, keys in enumerate(cite_tokens):
        text = text.replace(f"CITETOKEN{idx}END", r"\cite{" + keys + "}")
    return text


def make_architecture_figure() -> None:
    fig, ax = plt.subplots(figsize=(12.5, 3.4), dpi=220)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    boxes = [
        (0.02, "Simulink power stage\n+ fault injection", "#D9EAF7"),
        (0.215, "Synchronized signals\ncontroller + device", "#DFF2E1"),
        (0.41, "PWM-cycle aggregation\n+ event alignment", "#FFF1CC"),
        (0.605, "Physics features\n+ grouped validation", "#F6DFEC"),
        (0.80, "LR main classifier\n+ resistance threshold", "#E8E2F4"),
    ]
    for x, label, color in boxes:
        patch = FancyBboxPatch(
            (x, 0.32), 0.16, 0.38,
            boxstyle="round,pad=0.015,rounding_size=0.025",
            facecolor=color, edgecolor="#274C77", linewidth=1.4,
        )
        ax.add_patch(patch)
        ax.text(x + 0.08, 0.51, label, ha="center", va="center", fontsize=10.5, color="#17324D")
        if x < 0.79:
            ax.annotate("", xy=(x + 0.193, 0.51), xytext=(x + 0.165, 0.51),
                        arrowprops=dict(arrowstyle="->", color="#274C77", lw=1.6))
    ax.text(0.5, 0.15, "Run-level decisions + exact uncertainty + explicit claim boundary",
            ha="center", va="center", fontsize=10.5, color="#43576A", fontweight="bold")
    fig.tight_layout(pad=0.4)
    fig.savefig(FIG_DIR / "diagnostic_architecture.png", bbox_inches="tight", facecolor="white")
    plt.close(fig)


def copy_figures() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    sources = {
        ROOT / "ML" / "results" / "model_selection_notebook" / "pca_geometry.png": "pca_geometry.png",
        ROOT / "ML" / "results" / "ieee_computer_evidence_v17" / "figures" / "main_model_run_metrics.png": "main_model_run_metrics.png",
        ROOT / "ML" / "results" / "ieee_computer_evidence_v17" / "figures" / "high_r_detection_curve.png": "high_r_detection_curve.png",
        ROOT / "ML" / "results" / "ieee_computer_evidence_v17" / "figures" / "high_r_threshold_tradeoff.png": "high_r_threshold_tradeoff.png",
        ROOT / "ML" / "results" / "reference_hdf5_model_evaluation_v18" / "model_comparison.png": "synthetic_model_comparison.png",
    }
    for source, dest in sources.items():
        if not source.exists():
            raise FileNotFoundError(source)
        shutil.copy2(source, FIG_DIR / dest)
    make_architecture_figure()


def markdown_table(columns: list[str], rows: list[list[str]]) -> str:
    lines = ["| " + " | ".join(columns) + " |", "|" + "|".join(["---"] * len(columns)) + "|"]
    lines.extend("| " + " | ".join(row) + " |" for row in rows)
    return "\n".join(lines)


def build_markdown() -> str:
    out = [
        f"# {TITLE}",
        "",
        "**Mingyu Li; Qiang Wang†; Xiaojing Xiong†; Bianjia Wang†; Zhongxin Cheng*; Ying Jia***",
        "",
        AFFILIATION,
        "",
        "† Equal contribution and shared second authorship. * Corresponding authors.",
        "",
        AUTHOR_NOTE,
        "",
        f"> {DRAFT_NOTICE}",
        "",
        "## Abstract",
        "",
        ABSTRACT,
        "",
        f"**Index Terms:** {KEYWORDS}",
        "",
    ]
    for el in ELEMENTS:
        typ = el["type"]
        if typ == "h":
            marks = "#" * (el["level"] + 1)
            out.extend([f"{marks} {el['text']}", ""])
        elif typ == "p":
            out.extend([citation_word(el["text"]), ""])
        elif typ == "bullets":
            out.extend(["\n".join(f"- {citation_word(item)}" for item in el["items"]), ""])
        elif typ == "equation":
            out.extend([f"$${el['latex']} \tag{{{el['label'].strip('()')}}}$$", ""])
        elif typ == "table":
            out.extend([f"**{el['caption']}**", "", markdown_table(el["columns"], el["rows"]), ""])
            if el["note"]:
                out.extend([f"*Note: {el['note']}*", ""])
        elif typ == "figure":
            out.extend([f"![{el['caption']}](figures/{el['filename']})", "", f"*{el['caption']}*", ""])
            if el["source"]:
                out.extend([f"*Source note: {el['source']}*", ""])
        elif typ == "placeholder":
            out.extend([f"> **FIGURE PLACEHOLDER:** {el['instruction']}", "", f"*{el['caption']}*", ""])
    out.extend(["## REFERENCES", ""])
    out.extend(f"[{idx}] {text}" for idx, (_, text) in enumerate(REFERENCES, start=1))
    out.append("")
    return "\n".join(out)


def build_bibtex() -> str:
    return r"""@inproceedings{middlebrook1976,
  author={Middlebrook, R. D. and Cuk, Slobodan},
  title={A General Unified Approach to Modelling Switching-Converter Power Stages},
  booktitle={1976 IEEE Power Electronics Specialists Conference},
  pages={18--34}, year={1976}, doi={10.1109/PESC.1976.7072895}
}

@article{lee2019,
  author={Lee, Hyeon-Seok and Yun, Jae-Jung},
  title={High-Efficiency Bidirectional Buck-Boost Converter for Photovoltaic and Energy Storage Systems in a Smart Grid},
  journal={IEEE Transactions on Power Electronics}, volume={34}, number={5},
  pages={4316--4328}, year={2019}, doi={10.1109/TPEL.2018.2860059}
}

@inproceedings{kondrath2017,
  author={Kondrath, Nisha},
  title={Bidirectional DC-DC Converter Topologies and Control Strategies for Interfacing Energy Storage Systems in Microgrids: An Overview},
  booktitle={2017 5th IEEE International Conference on Smart Energy Grid Engineering},
  pages={341--345}, year={2017}
}

@inproceedings{sofla2011,
  author={Sofla, Mohammadhassan Abdollahi and Wang, Lingfeng},
  title={Control of DC-DC Bidirectional Converters for Interfacing Batteries in Microgrids},
  booktitle={2011 IEEE/PES Power Systems Conference and Exposition},
  pages={1--6}, year={2011}, doi={10.1109/PSCE.2011.5772602}
}

@article{xia2018,
  author={Xia, Kun and Hong, Xinwei and Yuan, Yin and Chen, Anghui},
  title={Bi-directional DC-DC Converter for Battery Energy Storage},
  journal={Journal of System Simulation}, volume={30}, number={8}, pages={3219--3228},
  year={2018}, doi={10.16182/j.issn1004731x.joss.201808049}, note={In Chinese}
}

@article{mahdavi2023,
  author={Mahdavi, Mohammad Saeed and Karimzadeh, Mohammad Saleh and Rahimi, Tohid and Gharehpetian, Gevork Babamalek},
  title={A Fault-Tolerant Bidirectional Converter for Battery Energy Storage Systems in DC Microgrids},
  journal={Electronics}, volume={12}, number={3}, pages={679}, year={2023}, doi={10.3390/electronics12030679}
}

@article{ding2024,
  author={Ding, Shichuan and Tang, Dewei and Hang, Jun and Zhao, Jifeng and Gui, Shuonan},
  title={Robust Open-Switch Fault Diagnosis of Bidirectional DC/DC Converters Based on Extended Kalman Filter With Multiple Corrections},
  journal={IEEE Transactions on Circuits and Systems I: Regular Papers}, volume={71}, number={9},
  pages={4363--4374}, year={2024}, doi={10.1109/TCSI.2024.3421655}
}

@article{alsheikh2015,
  author={Al-Sheikh, Hiba and Hoblos, Ghaleb and Moubayed, Nazih and Karami, Nabil},
  title={A Sensor Fault Diagnosis Scheme for a DC/DC Converter Used in Hybrid Electric Vehicles},
  journal={IFAC-PapersOnLine}, volume={48}, number={21}, pages={713--719},
  year={2015}, doi={10.1016/j.ifacol.2015.09.611}
}

@article{ouahabi2025,
  author={Ouahabi, Mohammed Said and Benyounes, Abdelhafid and Barkat, Said and Ihammouchen, Syphax and Rekioua, Toufik and Rabehi, Abdelaziz and El-Kenawy, El-Sayed M. and Alharbi, Amal H.},
  title={Real-Time Sensor Fault Tolerant Control of DC-DC Converters in DC Microgrids Using a Switching Unknown Input Observer},
  journal={IEEE Access}, volume={13}, pages={95838--95850}, year={2025}, doi={10.1109/ACCESS.2025.3571650}
}

@article{jin2024,
  author={Jin, Hailang and Gao, Zhiwei and Zuo, Zhiqiang and Zhang, Zhicheng and Wang, Yijing and Zhang, Aihua},
  title={A Combined Model-Based and Data-Driven Fault Diagnosis Scheme for Lithium-Ion Batteries},
  journal={IEEE Transactions on Industrial Electronics}, volume={71}, number={6},
  pages={6274--6284}, year={2024}, doi={10.1109/TIE.2023.3299029}
}

@article{liu2025,
  author={Liu, Kailong and Zhao, Shiwen and Wang, Yu and Li, Kang and Wang, Jiayue and Sun, Yaojie and Wu, Qiuwei and Peng, Qiao},
  title={Advanced Fault Diagnosis in Batteries: Insights into Fault Mechanisms, Sensor Fusion, and Artificial Intelligence},
  journal={Advances in Applied Energy}, volume={20}, pages={100247}, year={2025}, doi={10.1016/j.adapen.2025.100247}
}

@article{xu2025,
  author={Xu, Yiming and Ge, Xiaohua and Guo, Ruohan and Shen, Weixiang},
  title={Recent Advances in Model-Based Fault Diagnosis for Lithium-Ion Batteries: A Comprehensive Review},
  journal={Renewable and Sustainable Energy Reviews}, volume={207}, pages={114922}, year={2025}, doi={10.1016/j.rser.2024.114922}
}

@inproceedings{alsonisi2024,
  author={Alsonisi, Mohmaed and Ethni, Salaheddine and Elgendy, Mohammed and Ahmeid, Mohamed and Yildirim, Bortecene},
  title={DC-DC Bidirectional Converter for Battery Energy Storage System with Integrated Battery Management},
  booktitle={2024 IEEE International Conference and Exposition on Electric and Power Engineering},
  pages={89--93}, year={2024}, doi={10.1109/EPEi63510.2024.10758159}
}
"""


def table_latex(el: dict) -> str:
    cols = len(el["columns"])
    spec = "p{" + str(round(0.91 / cols, 3)) + r"\columnwidth}" * cols
    # IEEE columns are narrow; resize retains a legible relative layout in the template.
    header = " & ".join(latex_escape_plain(x) for x in el["columns"]) + r" \\ \hline"
    rows = [" & ".join(latex_escape_plain(x) for x in row) + r" \\" for row in el["rows"]]
    note = latex_escape_plain(el["note"]) if el["note"] else ""
    return (
        "\\begin{table}[!t]\n\\caption{" + latex_escape_plain(re.sub(r"^TABLE\s+[IVX]+\.\s*", "", el["caption"])) + "}\n"
        "\\centering\\scriptsize\n\\resizebox{\\columnwidth}{!}{%\n"
        "\\begin{tabular}{" + "l" * cols + "}\n\\hline\n" + header + "\n" + "\n".join(rows) +
        "\n\\hline\n\\end{tabular}}\n" +
        ("\\vspace{2pt}\\parbox{\\columnwidth}{\\scriptsize " + note + "}\n" if note else "") +
        "\\end{table}\n"
    )


def build_latex() -> str:
    lines = [
        r"\documentclass[conference]{IEEEtran}",
        r"\usepackage{graphicx}",
        r"\usepackage{amsmath,amssymb}",
        r"\usepackage{booktabs}",
        r"\usepackage{xcolor}",
        r"\usepackage{url}",
        r"\usepackage{cite}",
        r"\newcommand{\verify}[1]{\textcolor{red}{[VERIFY: #1]}}",
        "",
        r"\title{" + latex_escape_plain(TITLE) + "}",
        r"\author{",
        r"\IEEEauthorblockN{Mingyu Li, Qiang Wang$^{\dagger}$, Xiaojing Xiong$^{\dagger}$, Bianjia Wang$^{\dagger}$, Zhongxin Cheng$^{*}$, and Ying Jia$^{*}$}",
        r"\IEEEauthorblockA{[VERIFY: School/Department], The Chinese University of Hong Kong, Shenzhen, Shenzhen 518172, China\\",
        r"$^{\dagger}$Equal contribution and shared second authorship. $^{*}$Corresponding authors.\\",
        r"[VERIFY: insert corresponding-author e-mail addresses.]}",
        r"}",
        "",
        r"\begin{document}",
        r"\maketitle",
        r"\begin{abstract}",
        latex_escape_plain(ABSTRACT),
        r"\end{abstract}",
        r"\begin{IEEEkeywords}",
        latex_escape_plain(KEYWORDS),
        r"\end{IEEEkeywords}",
        r"\noindent\fbox{\parbox{0.96\columnwidth}{\footnotesize " + latex_escape_plain(DRAFT_NOTICE) + "}}",
        "",
    ]
    for el in ELEMENTS:
        typ = el["type"]
        if typ == "h":
            title = re.sub(r"^[IVX]+\.\s+", "", el["text"])
            if el["level"] == 1:
                lines.append(r"\section{" + latex_escape_plain(title) + "}")
            else:
                title = re.sub(r"^[A-Z]\.\s+", "", title)
                lines.append(r"\subsection{" + latex_escape_plain(title) + "}")
        elif typ == "p":
            lines.extend([latex_escape_plain(el["text"]), ""])
        elif typ == "bullets":
            lines.append(r"\begin{itemize}")
            lines.extend(r"\item " + latex_escape_plain(item) for item in el["items"])
            lines.extend([r"\end{itemize}", ""])
        elif typ == "equation":
            label = "eq:" + re.sub(r"\D", "", el["label"])
            lines.extend([r"\begin{equation}", el["latex"], r"\label{" + label + "}", r"\end{equation}"])
        elif typ == "table":
            lines.append(table_latex(el))
        elif typ == "figure":
            clean_caption = re.sub(r"^Fig\.\s+\d+\.\s*", "", el["caption"])
            lines.extend([
                r"\begin{figure}[!t]",
                r"\centering",
                r"\includegraphics[width=\columnwidth]{figures/" + el["filename"] + "}",
                r"\caption{" + latex_escape_plain(clean_caption) + "}",
                r"\end{figure}",
            ])
        elif typ == "placeholder":
            clean_caption = re.sub(r"^Fig\.\s+\d+\.\s*", "", el["caption"])
            lines.extend([
                r"\begin{figure}[!t]",
                r"\centering\fbox{\parbox{0.92\columnwidth}{\centering\vspace{0.22in}\footnotesize " + latex_escape_plain(el["instruction"]) + r"\vspace{0.22in}}}",
                r"\caption{" + latex_escape_plain(clean_caption) + "}",
                r"\end{figure}",
            ])
    lines.extend([
        r"\bibliographystyle{IEEEtran}",
        r"\bibliography{references}",
        r"\end{document}",
        "",
    ])
    return "\n".join(lines)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, width_twips=9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


VERIFY_RE = re.compile(r"(\[VERIFY:[^\]]+\])")


def add_marked_text(paragraph, text: str, bold: bool = False, italic: bool = False, size: float | None = None) -> None:
    for part in VERIFY_RE.split(citation_word(text)):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
        if part.startswith("[VERIFY:"):
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            run.font.color.rgb = RGBColor(156, 87, 0)


def configure_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.82)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(0.88)
    sec.right_margin = Inches(0.88)
    sec.header_distance = Inches(0.34)
    sec.footer_distance = Inches(0.34)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.2

    for name, size, color, before, after in (
        ("Heading 1", 14, "2E74B5", 15, 7),
        ("Heading 2", 11.5, "1F4D78", 10, 4),
    ):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    if "Caption Academic" not in [s.name for s in doc.styles]:
        st = doc.styles.add_style("Caption Academic", WD_STYLE_TYPE.PARAGRAPH)
        st.font.name = "Calibri"
        st.font.size = Pt(9)
        st.font.bold = True
        st.font.color.rgb = RGBColor(47, 74, 96)
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        st.paragraph_format.space_before = Pt(5)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.keep_with_next = True

    header = sec.header.paragraphs[0]
    header.text = "IEEE PAPER DRAFT  |  SIMULATION-ONLY  |  04 AUG 2026"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(92, 111, 126)
    footer = sec.footer.paragraphs[0]
    add_page_number(footer)
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(92, 111, 126)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(TITLE)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(22, 52, 77)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    add_marked_text(
        p,
        "Mingyu Li; Qiang Wang†; Xiaojing Xiong†; Bianjia Wang†; Zhongxin Cheng*; Ying Jia*",
        bold=True,
        size=11,
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    add_marked_text(p, AFFILIATION, italic=True, size=9.5)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    add_marked_text(p, AUTHOR_NOTE, size=8.5)

    box = doc.add_table(rows=1, cols=1)
    set_table_width(box)
    cell = box.cell(0, 0)
    set_cell_shading(cell, "EAF2F8")
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    add_marked_text(p, DRAFT_NOTICE, bold=True, size=8.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("Abstract-")
    r.bold = True
    add_marked_text(p, ABSTRACT)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Index Terms-")
    r.bold = True
    add_marked_text(p, KEYWORDS, italic=True)


def add_docx_table(doc: Document, el: dict) -> None:
    cap = doc.add_paragraph(style="Caption Academic")
    add_marked_text(cap, el["caption"], bold=True, size=9)
    table = doc.add_table(rows=1, cols=len(el["columns"]))
    table.style = "Table Grid"
    table.alignment = 1
    table.autofit = False
    set_table_width(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    prevent_row_split(header)
    widths = [9360 // len(el["columns"])] * len(el["columns"])
    for idx, (cell, text) in enumerate(zip(header.cells, el["columns"])):
        cell.width = Inches(widths[idx] / 1440)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "DCE6F1")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        add_marked_text(p, text, bold=True, size=8)
    for row_data in el["rows"]:
        row = table.add_row()
        prevent_row_split(row)
        for idx, (cell, text) in enumerate(zip(row.cells, row_data)):
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            add_marked_text(p, text, size=7.7)
    if el["note"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        add_marked_text(p, "Note: " + el["note"], italic=True, size=8)


def add_docx_figure(doc: Document, el: dict) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(FIG_DIR / el["filename"]), width=Inches(6.35))
    cap = doc.add_paragraph(style="Caption Academic")
    add_marked_text(cap, el["caption"], bold=True, size=9)
    if el["source"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        add_marked_text(p, "Evidence note: " + el["source"], italic=True, size=7.8)


def add_docx_placeholder(doc: Document, el: dict) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = 1
    set_table_width(table)
    prevent_row_split(table.rows[0])
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "F4F6F9")
    set_cell_margins(cell, top=380, start=240, bottom=380, end=240)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    add_marked_text(p, "FIGURE PLACEHOLDER\n" + el["instruction"], bold=True, size=9)
    cap = doc.add_paragraph(style="Caption Academic")
    add_marked_text(cap, el["caption"], bold=True, size=9)


def build_docx() -> None:
    doc = Document()
    configure_doc(doc)
    add_title_block(doc)
    for el in ELEMENTS:
        typ = el["type"]
        if typ == "h":
            p = doc.add_paragraph(style="Heading 1" if el["level"] == 1 else "Heading 2")
            add_marked_text(p, el["text"], bold=True)
        elif typ == "p":
            p = doc.add_paragraph()
            add_marked_text(p, el["text"])
        elif typ == "bullets":
            for item in el["items"]:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Inches(0.28)
                p.paragraph_format.first_line_indent = Inches(-0.18)
                p.paragraph_format.space_after = Pt(4)
                add_marked_text(p, item)
        elif typ == "equation":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(5)
            run = p.add_run(el["word"] + ("     " + el["label"] if el["label"] else ""))
            run.font.name = "Cambria Math"
            run.font.size = Pt(10)
            run.italic = True
        elif typ == "table":
            add_docx_table(doc, el)
        elif typ == "figure":
            add_docx_figure(doc, el)
        elif typ == "placeholder":
            add_docx_placeholder(doc, el)

    p = doc.add_paragraph(style="Heading 1")
    p.add_run("REFERENCES").bold = True
    for idx, (_, ref) in enumerate(REFERENCES, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.22)
        p.paragraph_format.space_after = Pt(3)
        add_marked_text(p, f"[{idx}] {ref}", size=8.5)

    props = doc.core_properties
    props.title = TITLE
    props.author = "; ".join(AUTHORS)
    props.subject = "Simulation-only IEEE paper sample and editable manuscript template"
    props.keywords = KEYWORDS
    props.comments = "Generated from verified local modeling results on 2026-08-04."
    doc.save(OUT / "paper_template_ieee.docx")


CHECKLIST = """# IEEE论文范文使用说明与投稿前核对清单

这套文件是基于当前已有仿真与计算结果写成的完整英文范文，不包含也不暗示硬件实验。正文中黄色或红色的 `[VERIFY: ...]` 是必须由作者确认的项目。

## 已经写实、可以保留的部分

- 主五分类数据：416个独立Run、28个工况、660个事件特征、6折工况分组OOF。
- LR/RF/Extra Trees的点估计均为Macro-F1=1.000、Accuracy=1.000；同时保留95%置信下界0.9912与健康误报率上界0.0253。
- 数据分布：26个常数特征、5803对高相关特征、PCA 90/95/99%维数10/21/53、有效秩9.758。
- 高阻负结果、v06直接器件电压/电流可观测性改造、10.5 mOhm阈值、12 mOhm压力测试边界。
- 合成HDF5的域迁移失败与seed/phase捷径风险；明确标为synthetic而非measured。

## 投稿前必须修改

1. 用最终Simulink模型截图替换Fig. 1占位框，并统一S1/S2、电流方向、Buck/Boost占空比定义。
2. 补全最终A_m、B_m或在附录给出对应模式方程；核对400 V、故障注入时刻、负载和电感/电容等参数。
3. 确认学院/部门英文全称、两位通讯作者邮箱、作者贡献；所有作者确认共同二作说明。
4. 向学校或课题组财务负责人确认基金正式英文名和项目号。若实际基金不是University Development Fund，替换整句，不要仅修改编号。
5. 给代码和数据设置公开仓库或机构存档地址；若暂不公开，按目标期刊政策改为“available on reasonable request”。
6. 在选定IEEE期刊/会议后按其页数、图分辨率、双盲和参考文献要求裁剪。当前LaTeX使用IEEEtran conference模板，Word稿用于便捷修改。
7. 最终检查模型结果图中的颜色、字号和术语是否与正文一致；不要把4320-run测量链投影称作动态功率级仿真或实验。
8. 如目标刊物要求算法原始文献，补充LR、RF、Extra Trees、XGBoost、MLP、KNN、1D-CNN及分组交叉验证的标准来源。当前参考文献暂按本地文件夹选择。

## 推荐的论文定位

建议定位为“一般IEEE会议或仿真/应用型期刊的simulation and computer-evidence paper”。核心贡献是Simulink故障建模、可观测性设计、工况分组验证和模型选择，而不是追求更复杂的神经网络。结论中继续保留“simulation-only”。

## 文件

- `paper_template_ieee.docx`：单栏可编辑Word范文，含图表和醒目标记。
- `paper_template_ieee.tex`：IEEEtran双栏LaTeX源文件。
- `references.bib`：从本地PDF核对的13篇参考文献。
- `paper_template_ieee_en.md`：便于全文检索和快速改写的Markdown源稿。
- `figures/`：本稿使用的现有结果图和诊断流程图。
"""


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    copy_figures()
    (OUT / "paper_template_ieee_en.md").write_text(build_markdown(), encoding="utf-8")
    (OUT / "paper_template_ieee.tex").write_text(build_latex(), encoding="utf-8")
    (OUT / "references.bib").write_text(build_bibtex(), encoding="utf-8")
    (OUT / "submission_checklist_cn.md").write_text(CHECKLIST, encoding="utf-8")
    build_docx()
    print(OUT)


if __name__ == "__main__":
    main()
